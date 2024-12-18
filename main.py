'''
Aşağıda programın çalışması için gerekli kütüphaneler
ve hemen yanlarında bu kütüphaneleri bilgisayarınıza
yüklemenizi sağlayan terminal promptları bulunuyor.
Eğer Python'u bilgisayarınıza  Anaconda Navigator ile
yüklediyseniz kodları Anaconda Prompt'a yazarak kütüphaneleri
yükleyebilirsiniz
'''
import requests  # pip install requests
from bs4 import BeautifulSoup  # pip install bs4
import pyodbc  # pip install pyodbc
import tkinter as tk # pip install tk
from tkinter import messagebox
from tkinter import filedialog
import webbrowser
import json
from docx import Document

# Bu fonksiyon oyunun ismini almamıza yarıyor
def get_game_name(app_id):
    url = f"https://store.steampowered.com/api/appdetails?appids={app_id}"
    response = requests.get(url)

    if response.status_code == 200:
        data = response.json()
        if data.get(str(app_id), {}).get('success', False):
            return data[str(app_id)]['data']['name']
    return "Böyle bir oyun yok..."

# Bu fonksiyon incelemeleri çekmeye yarıyor
def get_steam_reviews(app_id, num_pages=10):
    base_url = f"https://steamcommunity.com/app/{app_id}/reviews/"
    reviews = []

    for page in range(1, num_pages + 1):
        print(f"Sayfalar taranıyor: {page}...")
        url = f"{base_url}?p={page}&browsefilter=mostrecent"
        response = requests.get(url)

        if response.status_code != 200:
            print(f"Failed to fetch page {page}. Status code: {response.status_code}")
            break

        soup = BeautifulSoup(response.text, 'html.parser')
        review_blocks = soup.find_all('div', class_='apphub_CardTextContent')

        for block in review_blocks:
            review_text = block.get_text(strip=True)
            reviews.append(review_text)


    return reviews

# MSSQL Database'ine bağlantıyı kuran fonksiyon
def connect_to_mssql(server, database, username, password):
    conn = pyodbc.connect(
        f"DRIVER={{ODBC Driver 17 for SQL Server}};"
        f"SERVER={server};"
        f"DATABASE={database};"
        f"UID={username};"
        f"PWD={password}"
    )
    return conn

# Tabloları oluşturan fonksiyon
def setup_database(conn):
    cursor = conn.cursor()
    cursor.execute("""
        IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='Game' AND xtype='U')
        CREATE TABLE Game (
            AppId INT PRIMARY KEY,
            Name NVARCHAR(255) NOT NULL
        )
    """)
    cursor.execute("""
        IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='Review' AND xtype='U')
        CREATE TABLE Review (
            Id INT IDENTITY(1,1) PRIMARY KEY,
            AppId INT,
            Date DATE,
            Review NVARCHAR(MAX),
            FOREIGN KEY (AppId) REFERENCES Game(AppId)
        )
    """)
    conn.commit()

# Table'lara insertion yapan fonksiyon
def insert_data(conn, app_id, game_name, reviews):
    cursor = conn.cursor()
    cursor.execute("""
        IF NOT EXISTS (SELECT 1 FROM Game WHERE AppId = ?)
        INSERT INTO Game (AppId, Name) VALUES (?, ?)
    """, app_id, app_id, game_name)
    for review in reviews:
        cursor.execute("""
            INSERT INTO Review (AppId, Date, Review) VALUES (?, ?, ?)
        """, app_id, "2024-12-11", review)

    conn.commit()

# Word olarak kaydetme
def save_as_word():
    if not entry_app_id.get().isdigit():
        messagebox.showerror("Hata", "App ID bir sayı olmalı!")
        return

    global reviews, app_id, game_name

    app_id = entry_app_id.get().strip()
    game_name = get_game_name(app_id)
    reviews = get_steam_reviews(app_id)

    if not reviews:
        messagebox.showwarning("Uyarı", "Kaydedilecek bir inceleme bulunamadı")
        return

    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word files", "*.docx")],
        title="Save as Word"
    )
    if not file_path:
        return

    try:
        from docx import Document
        document = Document()
        document.add_heading(f"Reviews for '{game_name}' (App ID: {app_id})", level=1)

        for i, review in enumerate(reviews, start=1):
            document.add_paragraph(f"Review {i}:")
            document.add_paragraph(review)
            document.add_paragraph()  # Blank line between reviews

        document.save(file_path)
        messagebox.showinfo("Tamamlandı", f"İncelemeler word dosyası olarak kaydedildi:\n{file_path}")
    except Exception as e:
        messagebox.showerror("Hata", f"Kayıt başarısız: {e}")

# JSON olarak kaydetme
def save_as_json():
    if not entry_app_id.get().isdigit():
        messagebox.showerror("Hata", "App ID bir sayı olmalı!")
        return

    global reviews, app_id, game_name

    app_id = entry_app_id.get().strip()
    game_name = get_game_name(app_id)
    reviews = get_steam_reviews(app_id)

    if not reviews:
        messagebox.showwarning("Uyarı", "Kaydedilecek bir inceleme yok.")
        return

    file_path = filedialog.asksaveasfilename(
        defaultextension=".json",
        filetypes=[("JSON files", "*.json")],
        title="Save as JSON"
    )
    if not file_path:
        return

    data = {
        "App ID": app_id,
        "Game Name": game_name,
        "Reviews": reviews
    }

    try:
        with open(file_path, "w", encoding="utf-8") as json_file:
            json.dump(data, json_file, indent=4, ensure_ascii=False)
        messagebox.showinfo("Tamamlandı", f"İncelemeler JSON olarak kaydedildi: :\n{file_path}")
    except Exception as e:
        messagebox.showerror("Hata", f"Kayıt başarısız: {e}")


def fetch_and_store_reviews():
    app_id = entry_app_id.get().strip()
    if not app_id.isdigit():
        messagebox.showerror("Hata", "App ID bir sayı olmalı!")
        return

    #Buradan kaç sayfalık veri çekeceğinizi belirleyebilirsiniz.
    num_pages = 10
    game_name = get_game_name(app_id)
    if game_name == "No such game exists...":
        messagebox.showerror("Hata", "Geçersiz App ID! Oyun bulunamadı.")
        return

    reviews = get_steam_reviews(app_id, num_pages)
    if not reviews:
        messagebox.showwarning("Uyarı", f"Bu oyun için herhangi bir inceleme bulunamadı: '{game_name}'.")
        return


    global current_reviews, current_game_name, current_app_id
    current_reviews = reviews
    current_game_name = game_name
    current_app_id = app_id


    # MSSQL bağlantısı ve insertion
    try:
        conn = connect_to_mssql(server, database, username, password)
        setup_database(conn)
        insert_data(conn, app_id, game_name, reviews)
        conn.close()
        messagebox.showinfo("Tamamlandı",  f'{game_name}'" için " f"{len(reviews)} inceleme çekildi.")
    except Exception as e:
        messagebox.showerror("Database Hatası", f"Error: {e}")


def open_steamdb_link():
    webbrowser.open("https://steamdb.info/apps/")

#Kullanıcı arayüzü, Ana window
window = tk.Tk()
window.title("Steam Review Scraper")
window.geometry("400x250")


label_instruction = tk.Label(window, text="Oyunun App ID'sini giriniz: ")
label_instruction.pack(pady=5)

entry_app_id = tk.Entry(window, width=30)
entry_app_id.pack(pady=5)

button_frame = tk.Frame(window)
button_frame.pack(pady=10)

button_save_word = tk.Button(button_frame, text="Word olarak kaydet", command=save_as_word)
button_save_word.pack(side=tk.LEFT, padx=5)

button_save_json = tk.Button(button_frame, text="JSON olarak kaydet", command=save_as_json)
button_save_json.pack(side=tk.LEFT, padx=5)

# Insert to Database button
button_fetch = tk.Button(button_frame, text="Veritabanına kaydet", command=fetch_and_store_reviews)
button_fetch.pack(side=tk.LEFT, padx=5)

label_link = tk.Label(window, text="Aradığınız Oyunun App ID'sini öğrenmek için tıklayın.", fg="blue", cursor="hand2")
label_link.pack(pady=10)
label_link.bind("<Button-1>", lambda e: open_steamdb_link())

'''
Database creditentials
Aşağıya bilgisayarınızda oluşturduğunuz database server'ının 
ismini, database in ismini ve kullanıcı bilgilerinizi girmeniz 
gerekiyor.  
'''
server = ""
database = ""
username = ""
password = ""

window.mainloop()
